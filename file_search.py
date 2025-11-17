import os
import re
import sys
import subprocess
import threading
import tkinter as tk
from tkinter import filedialog, scrolledtext, messagebox, ttk
from PyPDF2 import PdfReader
from docx import Document

# Conditional imports for document types
try:
    from openpyxl import load_workbook
    EXCEL_XLSX_SUPPORT = True
except ImportError:
    EXCEL_XLSX_SUPPORT = False

try:
    import xlrd
    EXCEL_XLS_SUPPORT = True
except ImportError:
    EXCEL_XLS_SUPPORT = False

try:
    import win32com.client
    WIN32COM_SUPPORT = True
except ImportError:
    WIN32COM_SUPPORT = False

# Constants
SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".xlsx", ".xls")
WINDOW_TITLE = "Modern Offline Document Keyword Searcher"
WINDOW_SIZE = "950x750" # Slightly larger for modern feel
THEME_NAME = "clam" # Or "alt", "default", "vista", "xpnative" (Windows)

# --- Helper Functions (No Change Needed) ---
def _number_to_excel_column(n):
    """Convert a column number (1-based) to Excel column letter (A, B, ..., Z, AA, AB, ...)."""
    result = ""
    while n > 0:
        n -= 1
        result = chr(65 + (n % 26)) + result
        n //= 26
    return result

# --- Search Logic (Updated for XLS support and minor refactoring) ---
def search_document_for_keywords(filepath, keywords, log_callback):
    """
    Search a single document (PDF, DOCX, or Excel) for specified keywords.
    
    Args:
        filepath: Path to the document file
        keywords: List of keywords to search for
        log_callback: Function to call for logging messages
        
    Returns:
        List of page numbers for PDFs, sheet names for Excel, or ["Document Content"] for DOCX if found
    """
    found_info = []
    # Using word boundaries for more precise matching
    keyword_pattern = re.compile(
        r'\b(?:' + '|'.join(re.escape(k) for k in keywords) + r')\b',
        re.IGNORECASE
    )
    filename = os.path.basename(filepath)
    file_extension = filename.lower()

    if file_extension.endswith(".pdf"):
        try:
            with open(filepath, 'rb') as file:
                reader = PdfReader(file)
                for page_num, page in enumerate(reader.pages, start=1):
                    text = page.extract_text()
                    if text and keyword_pattern.search(text):
                        found_info.append(page_num)
        except Exception as e:
            log_callback(f"Error processing PDF '{filename}': {e}")

    elif file_extension.endswith(".docx"):
        try:
            doc = Document(filepath)
            doc_text = []
            for para in doc.paragraphs:
                doc_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        doc_text.append(cell.text)
            full_text = "\n".join(doc_text)
            if keyword_pattern.search(full_text):
                found_info.append("Document Content")
        except Exception as e:
            log_callback(f"Error processing DOCX '{filename}': {e}")

    elif file_extension.endswith(".xlsx"):
        if not EXCEL_XLSX_SUPPORT:
            log_callback(f"Excel (.xlsx) support not available. Install openpyxl: pip install openpyxl")
            return found_info
        
        try:
            workbook = load_workbook(filepath, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                for row in sheet.iter_rows(values_only=False):
                    for cell in row:
                        if cell.value is not None:
                            cell_text = str(cell.value)
                            if keyword_pattern.search(cell_text):
                                row_idx = cell.row
                                col_idx = cell.column
                                col_letter = _number_to_excel_column(col_idx)
                                cell_ref = f"{col_letter}{row_idx}"
                                found_info.append({
                                    'type': 'excel_cell',
                                    'file_type': '.xlsx',
                                    'sheet': sheet_name,
                                    'cell': cell_ref,
                                    'row': row_idx,
                                    'col': col_idx
                                })
        except Exception as e:
            log_callback(f"Error processing Excel (.xlsx) file '{filename}': {e}")

    elif file_extension.endswith(".xls"):
        if not EXCEL_XLS_SUPPORT:
            log_callback(f"Excel (.xls) support not available. Install xlrd: pip install xlrd")
            return found_info
        
        try:
            workbook = xlrd.open_workbook(filepath)
            for sh_idx in range(workbook.nsheets):
                sheet = workbook.sheet_by_index(sh_idx)
                sheet_name = sheet.name
                for row_idx in range(sheet.nrows):
                    for col_idx in range(sheet.ncols):
                        cell_value = sheet.cell_value(row_idx, col_idx)
                        if cell_value is not None and str(cell_value).strip() != "":
                            cell_text = str(cell_value)
                            if keyword_pattern.search(cell_text):
                                col_letter = _number_to_excel_column(col_idx + 1) # xlrd is 0-indexed, _number_to_excel_column is 1-indexed
                                cell_ref = f"{col_letter}{row_idx + 1}" # xlrd is 0-indexed for rows too
                                found_info.append({
                                    'type': 'excel_cell',
                                    'file_type': '.xls',
                                    'sheet': sheet_name,
                                    'cell': cell_ref,
                                    'row': row_idx + 1,
                                    'col': col_idx + 1
                                })
        except Exception as e:
            log_callback(f"Error processing Excel (.xls) file '{filename}': {e}")

    return found_info


def search_files_in_folder_gui(folder_path, keywords, log_callback, progress_callback=None):
    """
    Search for keywords in all supported documents within a folder.
    
    Args:
        folder_path: Path to the folder to search
        keywords: List of keywords to search for
        log_callback: Function to call for logging messages
        progress_callback: Function to call to update progress bar (current, total)
        
    Returns:
        Dictionary mapping filenames to their search results
    """
    found_keywords_in_docs = {}
    if not keywords:
        log_callback("No keywords provided for search.")
        return {}

    if not os.path.isdir(folder_path):
        log_callback(f"Error: Folder '{folder_path}' not found.")
        return {}

    all_files_in_folder = [
        f for f in os.listdir(folder_path)
        if f.lower().endswith(SUPPORTED_EXTENSIONS) and os.path.isfile(os.path.join(folder_path, f))
    ]

    if not all_files_in_folder:
        log_callback(f"No supported documents (PDF, DOCX, Excel) found in '{folder_path}'.")
        return {}
    
    total_files = len(all_files_in_folder)
    for i, filename in enumerate(all_files_in_folder, start=1):
        log_callback(f"Processing '{filename}' ({i}/{total_files})...")
        if progress_callback:
            progress_callback(i, total_files) # Update progress with current and total
        filepath = os.path.join(folder_path, filename)
        found_details = search_document_for_keywords(filepath, keywords, log_callback)
        if found_details:
            found_keywords_in_docs[filename] = found_details

    return found_keywords_in_docs


# --- GUI Class (Extensively Updated) ---
class DocumentSearchGUI:
    """GUI application for searching keywords in PDF, DOCX, and Excel documents."""
    
    def __init__(self, master):
        self.master = master
        master.title(WINDOW_TITLE)
        master.geometry(WINDOW_SIZE)
        self.doc_folder = os.getcwd()

        self._apply_style()
        self._create_frames()
        self._create_controls()
        self._create_progress_area()
        self._create_results_area()

    def _apply_style(self):
        """Apply modern ttk styling."""
        style = ttk.Style(self.master)
        style.theme_use(THEME_NAME)

        # Configure general styles
        style.configure('TFrame', background='#e0e0e0')
        style.configure('TLabel', background='#e0e0e0', font=('Segoe UI', 10))
        style.configure('TButton', font=('Segoe UI', 10, 'bold'), padding=6)
        style.map('TButton', background=[('active', '#c0c0c0')])
        style.configure('TEntry', font=('Segoe UI', 10), padding=5)
        style.configure('TScrolledText', font=('Segoe UI', 10)) # Custom widget, so tags are more important

        # Custom styles for specific elements
        style.configure('Watermark.TLabel', foreground='#9aa0a6', font=('Segoe UI', 9, 'italic'))
        style.configure('Header.TLabel', font=('Segoe UI', 11, 'bold'))
        style.configure('Results.TLabel', font=('Segoe UI', 10, 'bold'))


    def _create_frames(self):
        """Create and pack all frame widgets using ttk.Frame for consistency."""
        self.main_frame = ttk.Frame(self.master, padding="15 15 15 15")
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        self.controls_frame = ttk.Frame(self.main_frame, padding="10 10 10 10", relief=tk.GROOVE)
        self.controls_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 10))

        self.progress_frame = ttk.Frame(self.main_frame, padding="5 5 5 5")
        self.progress_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 5))

        self.results_frame = ttk.Frame(self.main_frame, padding="10 10 10 10", relief=tk.GROOVE)
        self.results_frame.pack(side=tk.BOTTOM, fill=tk.BOTH, expand=True, pady=(5, 0))

        self.watermark_frame = ttk.Frame(self.main_frame, padding="0 5 0 0")
        self.watermark_frame.pack(side=tk.BOTTOM, fill=tk.X)
        self.watermark_label = ttk.Label(
            self.watermark_frame,
            text="Created by Vinay Bhaskarla",
            style='Watermark.TLabel'
        )
        self.watermark_label.pack(pady=(0, 4))


    def _create_controls(self):
        """Create control widgets (folder selection, keywords, buttons)."""
        self.controls_frame.grid_columnconfigure(0, weight=0)
        self.controls_frame.grid_columnconfigure(1, weight=1)
        self.controls_frame.grid_columnconfigure(2, weight=0) # Make button column non-expanding

        ttk.Label(self.controls_frame, text="Document Directory:", style='Header.TLabel').grid(row=0, column=0, sticky=tk.W, pady=5, padx=(0,10))

        self.folder_path_var = tk.StringVar(value=self.doc_folder)
        self.folder_entry = ttk.Entry(self.controls_frame, textvariable=self.folder_path_var)
        self.folder_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")

        self.browse_button = ttk.Button(self.controls_frame, text="Browse", command=self.browse_folder)
        self.browse_button.grid(row=0, column=2, padx=5, pady=5, sticky="e")

        ttk.Label(self.controls_frame, text="Keywords (comma-separated):", style='Header.TLabel').grid(row=1, column=0, sticky=tk.W, pady=5, padx=(0,10))
        self.keyword_entry = ttk.Entry(self.controls_frame)
        self.keyword_entry.grid(row=1, column=1, padx=5, pady=5, sticky="ew")
        self.keyword_entry.bind("<Return>", self.start_search_thread)

        self.search_button = ttk.Button(self.controls_frame, text="Search Documents", command=self.start_search_thread)
        self.search_button.grid(row=1, column=2, padx=5, pady=5, sticky="e")
        
        # Add Clear button for keywords
        self.clear_keywords_button = ttk.Button(self.controls_frame, text="Clear Keywords", command=lambda: self.keyword_entry.delete(0, tk.END))
        self.clear_keywords_button.grid(row=2, column=1, padx=5, pady=5, sticky="e")


    def _create_progress_area(self):
        """Create and configure the progress bar and status label."""
        self.progress_label = ttk.Label(self.progress_frame, text="Ready to search.")
        self.progress_label.pack(fill=tk.X, pady=(0, 5))

        self.progress_bar = ttk.Progressbar(self.progress_frame, mode='determinate', length=800)
        self.progress_bar.pack(fill=tk.X, pady=5)
        self.progress_bar.stop()
        self.progress_bar['value'] = 0


    def _create_results_area(self):
        """Create and configure the results text area."""
        ttk.Label(self.results_frame, text="Search Results:", style='Results.TLabel').pack(pady=(0, 5), anchor=tk.W)

        self.results_text = scrolledtext.ScrolledText(
            self.results_frame,
            wrap=tk.WORD,
            state=tk.DISABLED,
            width=80,
            height=12,
            font=('Segoe UI', 10),
            relief=tk.FLAT, # Flat relief for a modern look
            background='#f9f9f9', # Slightly off-white background
            foreground='#333333' # Darker text color
        )
        self.results_text.pack(fill=tk.BOTH, expand=True)

        # Modernized tag configurations
        self.results_text.tag_config("link", foreground="#007bff", underline=True, font=('Segoe UI', 10, 'underline')) # Bootstrap primary blue
        self.results_text.tag_bind("link", "<Enter>", lambda e: e.widget.config(cursor="hand2"))
        self.results_text.tag_bind("link", "<Leave>", lambda e: e.widget.config(cursor=""))
        self.results_text.tag_config("log", foreground="#6c757d", font=('Segoe UI', 9, 'italic')) # Bootstrap secondary gray
        self.results_text.tag_config("file_header", foreground="#343a40", font=('Segoe UI', 10, 'bold')) # Darker header
        self.results_text.tag_config("sheet_header", foreground="#495057", font=('Segoe UI', 10, 'bold')) # Slightly lighter header


    def browse_folder(self):
        """Open a dialog for the user to select a directory."""
        folder_selected = filedialog.askdirectory(initialdir=self.doc_folder)
        if folder_selected:
            self.doc_folder = folder_selected
            self.folder_path_var.set(self.doc_folder)

    def log_message(self, message, tag="log"):
        """Thread-safe way to insert log messages into the results_text widget."""
        self.master.after(0, self._insert_log, message, tag)

    def _insert_log(self, message, tag):
        """Insert a log message into the results text widget."""
        self.results_text.config(state=tk.NORMAL)
        self.results_text.insert(tk.END, f"{message}\n", tag)
        self.results_text.see(tk.END)
        self.results_text.config(state=tk.DISABLED)

    def update_progress(self, current, total):
        """Thread-safe way to update the progress bar and label."""
        self.master.after(0, self._update_progress_gui, current, total)

    def _update_progress_gui(self, current, total):
        """Update the progress bar and label on the GUI thread."""
        if total > 0:
            percentage = (current / total) * 100
            self.progress_bar['value'] = percentage
            self.progress_label.config(text=f"Processing file {current} of {total} ({percentage:.0f}%)...")
        else:
            self.progress_bar['value'] = 0
            self.progress_label.config(text="Starting search...")


    def start_search_thread(self, event=None):
        """Start the document search in a new thread to keep the GUI responsive."""
        keywords_input = self.keyword_entry.get().strip()
        search_keywords = [kw.strip() for kw in keywords_input.split(',') if kw.strip()]

        current_doc_folder = self.folder_path_var.get()
        if not os.path.exists(current_doc_folder):
            messagebox.showerror("Error", f"The specified document directory does not exist:\n{current_doc_folder}")
            return

        self.doc_folder = current_doc_folder

        if not search_keywords:
            messagebox.showwarning("Input Error", "Please enter keywords to search for.")
            return

        self.results_text.config(state=tk.NORMAL)
        self.results_text.delete(1.0, tk.END)
        self.results_text.config(state=tk.DISABLED)

        self.log_message(f"Starting search for '{', '.join(search_keywords)}' in '{self.doc_folder}'...")
        self.search_button.config(state=tk.DISABLED)
        self.browse_button.config(state=tk.DISABLED)
        self.keyword_entry.config(state=tk.DISABLED)
        self.clear_keywords_button.config(state=tk.DISABLED)

        self.progress_bar.config(mode='determinate') # Ensure it's determinate for file processing
        self.progress_bar.start() # Start with indeterminate movement until total files are known
        self.progress_bar['value'] = 0
        self.progress_label.config(text="Discovering files...")


        search_thread = threading.Thread(
            target=self.perform_search_threaded,
            args=(self.doc_folder, search_keywords)
        )
        search_thread.daemon = True
        search_thread.start()

    def perform_search_threaded(self, folder_path, search_keywords):
        """Run the search logic and update the GUI upon completion."""
        results = search_files_in_folder_gui(folder_path, search_keywords, self.log_message, self.update_progress)
        self.master.after(0, self.display_results, results, search_keywords)

    def display_results(self, results, search_keywords):
        """Display the search results in the GUI."""
        self.progress_bar.stop()
        self.progress_bar['value'] = 0 # Reset progress bar
        self.progress_label.config(text="Search complete.")

        self.search_button.config(state=tk.NORMAL)
        self.browse_button.config(state=tk.NORMAL)
        self.keyword_entry.config(state=tk.NORMAL)
        self.clear_keywords_button.config(state=tk.NORMAL)


        self.results_text.config(state=tk.NORMAL)
        self.results_text.insert(tk.END, "\n--- Search Complete ---\n\n", "log")

        if results:
            self.results_text.insert(tk.END, "Found keywords in the following documents:\n", "log")
            for doc_file, details in results.items():
                full_path = os.path.join(self.doc_folder, doc_file)
                self.results_text.insert(tk.END, f"\nFile: {doc_file}\n", "file_header")

                file_extension = doc_file.lower()

                if file_extension.endswith(".pdf"):
                    self.results_text.insert(tk.END, "  Keywords found on pages: ")
                    for i, page_num in enumerate(details):
                        link_text = f"Page {page_num}"
                        unique_tag = f"link_{full_path}_{page_num}"
                        self.results_text.insert(tk.END, link_text, ("link", unique_tag))
                        self.results_text.tag_bind(
                            unique_tag, "<Button-1>",
                            lambda e, p=full_path, pg=page_num: self.open_document(e, p, pg)
                        )
                        if i < len(details) - 1:
                            self.results_text.insert(tk.END, ", ")
                    self.results_text.insert(tk.END, "\n")

                elif file_extension.endswith(".docx"):
                    self.results_text.insert(tk.END, "  Keywords found: ")
                    link_text = "Click to Open Document"
                    unique_tag = f"link_{full_path}_docx"
                    self.results_text.insert(tk.END, link_text, ("link", unique_tag))
                    self.results_text.tag_bind(
                        unique_tag, "<Button-1>",
                        lambda e, p=full_path, pg=None: self.open_document(e, p, pg)
                    )
                    self.results_text.insert(tk.END, "\n")

                elif file_extension.endswith((".xlsx", ".xls")):
                    # Group cells by sheet
                    sheets_dict = {}
                    for cell_info in details:
                        if isinstance(cell_info, dict) and cell_info.get('type') == 'excel_cell':
                            sheet_name = cell_info['sheet']
                            if sheet_name not in sheets_dict:
                                sheets_dict[sheet_name] = []
                            sheets_dict[sheet_name].append(cell_info)
                    
                    for sheet_name, cells in sheets_dict.items():
                        self.results_text.insert(tk.END, f"  Sheet '{sheet_name}' - Cells: ", "sheet_header")
                        unique_cells = [] # To avoid duplicate cell references if multiple keywords found in same cell
                        for cell_info in cells:
                            if cell_info['cell'] not in unique_cells:
                                unique_cells.append(cell_info['cell'])
                        
                        for i, cell_ref in enumerate(unique_cells):
                            link_text = cell_ref
                            unique_tag = f"link_{full_path}_{sheet_name}_{cell_ref}"
                            self.results_text.insert(tk.END, link_text, ("link", unique_tag))
                            self.results_text.tag_bind(
                                unique_tag, "<Button-1>",
                                lambda e, p=full_path, s=sheet_name, c=cell_ref: 
                                self.open_excel_at_cell(e, p, s, c)
                            )
                            if i < len(unique_cells) - 1:
                                self.results_text.insert(tk.END, ", ")
                        self.results_text.insert(tk.END, "\n")
        else:
            self.results_text.insert(
                tk.END,
                "\nNo supported documents found containing the specified keywords in the selected directory.\n",
                "log"
            )

        self.results_text.see(tk.END)
        self.results_text.config(state=tk.DISABLED)

    def _get_program_paths(self, program_name, common_paths):
        """Helper to find common program executable paths."""
        paths = []
        if sys.platform == "win32":
            # Check program files (x86) and program files
            for pf in [os.environ.get('ProgramFiles(x86)'), os.environ.get('ProgramFiles')]:
                if pf:
                    for common_path in common_paths:
                        full_path = os.path.join(pf, common_path, program_name)
                        if os.path.exists(full_path):
                            paths.append(full_path)
            # Check PATH environment variable
            for path_dir in os.environ["PATH"].split(os.pathsep):
                exe_path = os.path.join(path_dir, program_name)
                if os.path.exists(exe_path):
                    paths.append(exe_path)
        return list(set(paths)) # Return unique paths


    def open_document(self, event, doc_path, page_num=None):
        """
        Open a document (PDF, DOCX, or Excel) using the default system viewer.
        
        Args:
            event: Tkinter event object
            doc_path: Path to the document file
            page_num: Optional page number for PDF files (not used for DOCX/Excel)
        """
        try:
            absolute_doc_path = os.path.abspath(doc_path)
            platform = sys.platform
            
            if platform == "win32":
                if page_num is not None and doc_path.lower().endswith('.pdf'):
                    success = False
                    
                    # More robust Adobe Reader paths
                    adobe_reader_exes = self._get_program_paths(
                        "AcroRd32.exe", 
                        [r'Adobe\Acrobat Reader DC\Reader', r'Adobe\Acrobat Reader\Reader', r'Adobe\Reader']
                    ) + self._get_program_paths(
                        "Acrobat.exe",
                        [r'Adobe\Acrobat DC\Acrobat', r'Adobe\Acrobat\Acrobat']
                    )

                    for adobe_exe in adobe_reader_exes:
                        if os.path.exists(adobe_exe):
                            try:
                                subprocess.Popen([adobe_exe, '/A', f'page={page_num}', absolute_doc_path], shell=False)
                                success = True
                                break
                            except Exception:
                                continue
                    
                    if not success:
                        try:
                            # Fallback 1: Try file URI with page number (works with some viewers like modern Edge/Chrome)
                            uri_path = absolute_doc_path.replace(os.sep, '/')
                            full_uri = f'file:///{uri_path}#page={page_num}'
                            subprocess.Popen(f'start "" "{full_uri}"', shell=True)
                            success = True
                        except Exception:
                            pass
                    
                    if not success:
                        # Fallback 2: Open without page number
                        subprocess.Popen(f'start "" "{absolute_doc_path}"', shell=True)
                else:
                    subprocess.Popen(f'start "" "{absolute_doc_path}"', shell=True)
            elif platform == "darwin":
                if page_num is not None:
                    # macOS Preview supports #page parameter
                    subprocess.Popen(['open', f'{absolute_doc_path}#page={page_num}'])
                else:
                    subprocess.Popen(['open', absolute_doc_path])
            else:
                if page_num is not None:
                    # Linux xdg-open sometimes supports #page, but not universally
                    subprocess.Popen(['xdg-open', f'{absolute_doc_path}#page={page_num}'])
                else:
                    subprocess.Popen(['xdg-open', absolute_doc_path])
        except Exception as e:
            messagebox.showerror(
                "Error",
                f"Could not open document.\nError: {e}\n\n"
                "Please ensure you have a default viewer installed for this document type "
                "and that the file path is valid."
                f"\nAttempted path: {absolute_doc_path}{f'#page={page_num}' if page_num else ''}"
            )

    def open_excel_at_cell(self, event, excel_path, sheet_name, cell_ref):
        """
        Open Excel file and navigate to a specific cell.
        
        Args:
            event: Tkinter event object
            excel_path: Path to the Excel file
            sheet_name: Name of the sheet
            cell_ref: Cell reference (e.g., "A1", "B5")
        """
        try:
            absolute_excel_path = os.path.abspath(excel_path)
            platform = sys.platform
            
            if platform == "win32":
                if WIN32COM_SUPPORT:
                    try:
                        excel = win32com.client.Dispatch("Excel.Application")
                        excel.Visible = True
                        workbook = excel.Workbooks.Open(absolute_excel_path)
                        
                        # Ensure the correct sheet is activated
                        # win32com.client uses 1-based indexing for Worksheets
                        # or can find by name. Find by name is more robust.
                        found_sheet = False
                        for ws in workbook.Sheets:
                            if ws.Name == sheet_name:
                                ws.Activate()
                                found_sheet = True
                                break
                        
                        if found_sheet:
                            # Select the range and scroll to it
                            excel.Application.Wait(excel.Application.Now + (1/24/60/60*2)) # Small delay
                            worksheet = workbook.ActiveSheet
                            worksheet.Range(cell_ref).Select()
                            excel.Application.ActiveWindow.ScrollRow = worksheet.Range(cell_ref).Row
                            excel.Application.ActiveWindow.ScrollColumn = worksheet.Range(cell_ref).Column
                            return # Successfully opened and navigated via COM
                        else:
                            # If sheet not found by name, still open the workbook
                            messagebox.showwarning(
                                "Excel Navigation Warning",
                                f"Sheet '{sheet_name}' not found in Excel file.\n\n"
                                f"The Excel file has been opened, but you may need to manually navigate to:\n"
                                f"Sheet: {sheet_name}\n"
                                f"Cell: {cell_ref}"
                            )
                            return # Still count as a successful open, but with warning
                    except Exception as com_e:
                        # Log COM error, but don't stop, try other methods
                        self.log_message(f"COM automation failed: {com_e}", "log")
                        pass # Fall through to other methods if COM fails
                
                # Fallback: Try common Excel executable paths
                excel_exes = self._get_program_paths(
                    "EXCEL.EXE",
                    [r'Microsoft Office\root\Office16', r'Microsoft Office\Office16',
                     r'Microsoft Office\root\Office15', r'Microsoft Office\Office15'] # Add more versions
                )
                for excel_exe in excel_exes:
                    if os.path.exists(excel_exe):
                        try:
                            subprocess.Popen([excel_exe, absolute_excel_path], shell=False)
                            messagebox.showinfo(
                                "Excel Opened",
                                f"Excel file opened.\n\n"
                                f"Please navigate to:\n"
                                f"Sheet: {sheet_name}\n"
                                f"Cell: {cell_ref}"
                            )
                            return
                        except Exception:
                            continue
                
                # Final fallback: open with default application (via shell)
                subprocess.Popen(f'start "" "{absolute_excel_path}"', shell=True)
                messagebox.showinfo(
                    "Excel Opened",
                    f"Excel file opened using default application.\n\n"
                    f"Please navigate to:\n"
                    f"Sheet: {sheet_name}\n"
                    f"Cell: {cell_ref}"
                )
            elif platform == "darwin":
                # macOS - try to open with Microsoft Excel specifically
                try:
                    subprocess.Popen(['open', '-a', 'Microsoft Excel', absolute_excel_path])
                except Exception:
                    # Fallback to default application
                    subprocess.Popen(['open', absolute_excel_path])
                messagebox.showinfo(
                    "Excel Opened",
                    f"Excel file opened.\n\n"
                    f"Please navigate to:\n"
                    f"Sheet: {sheet_name}\n"
                    f"Cell: {cell_ref}"
                )
            else:
                # Linux - open with default application
                subprocess.Popen(['xdg-open', absolute_excel_path])
                messagebox.showinfo(
                    "Excel Opened",
                    f"Excel file opened.\n\n"
                    f"Please navigate to:\n"
                    f"Sheet: {sheet_name}\n"
                    f"Cell: {cell_ref}"
                )
        except Exception as e:
            messagebox.showerror(
                "Error",
                f"Could not open Excel file.\nError: {e}\n\n"
                f"File: {absolute_excel_path}\n"
                f"Sheet: {sheet_name}\n"
                f"Cell: {cell_ref}"
            )


if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentSearchGUI(root)
    root.mainloop()